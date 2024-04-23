#%%
import json
import zipfile
from datetime import datetime
from typing import Dict, Optional

from docx import Document


#%%
def parse_client_info(doc:Document) -> Optional[Dict[str,str]]:
    if len(doc.tables) == 0:
        return None 
    
    # initialize values used to parse the first table in the document
    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # get list of the unique strings from each cell in the table
    list_info = []
    for r in range(num_rows):
        for c in range(num_cols):
            # parse out the text for the current cell
            text = table.cell(r,c).text.strip()
            text = text.replace('\t', "")
            text = text.replace('\n', "")
            text = text.replace('\u2019', "'")
            text = text.replace('\u2013', "-")

            # ensure the text is not empty
            item_text = text.strip()
            if len(item_text) > 0:
                # if the string contains a comma, only take the first part of it
                if ',' in item_text:
                    item_text = item_text.split(',')[0]

                # ensure the string is not a duplicate
                if item_text not in list_info:
                    list_info.append(item_text)

    # loop through the list of strings
    client_info: Dict[str,str] = {}
    for item in list_info:
        # split the string on the colon and get the key and value
        item_split = item.split(':', 1)
        item_key = item_split[0].strip()
        item_value = item_split[1].strip()
        #print(item_value)

        if item_key.startswith('Date'):
            # split date into month, day and year
            date_info = item_value.split('-')
            date_value = datetime(int(date_info[2]), int(date_info[0]), int(date_info[1]))
            item_value = datetime.strftime(date_value, "%Y-%m-%d")

        client_info[item_key] = item_value

    return client_info


#%%
zip_filename =  "../data/assessments/clients-20240410.zip"
with zipfile.ZipFile(zip_filename, 'r') as zf:
    for f in zf.namelist():
        if f.endswith('.docx'):
            with zf.open(f, 'r') as fp:
                document = Document(fp)
                info = parse_client_info(document)
                print(info)


#%%
