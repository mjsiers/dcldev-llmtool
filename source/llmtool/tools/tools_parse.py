import logging
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

from ..data.models import DocumentSchema
from .tools_embed import embed_text

# configure logging
logger = logging.getLogger(__name__)


def is_section_text(text: str, sections_text: List[str]) -> int:
    for idx, item in enumerate(sections_text):
        if text.startswith(item):
            return idx
    return -1


def get_age(date: str, birth_date: str) -> Optional[float]:
    # ensure the date and birth date are valid
    if len(date) == 0 or len(birth_date) == 0:
        return None

    # parse the date values
    today = datetime.strptime(date, "%Y-%m-%d")
    dob = datetime.strptime(birth_date, "%Y-%m-%d")

    # compute the age value
    delta_days = (today - dob).days
    age_value = delta_days / 365.25
    return age_value


def get_grade(grade: str) -> Optional[int]:
    # ensure the grade is not empty
    grade = grade.strip()
    if len(grade) == 0:
        return None

    # determine the number of numeric characters in the grade
    num_length = 0
    for c in grade:
        if not c.isdigit():
            break
        num_length += 1

    # ensure we have some numeric characters
    if num_length == 0:
        return None

    # convert the grade to a numeric value
    grade = grade[:num_length]
    return int(grade)


def parse_key_reasons(dict_keywords: Dict[str, int], list_reasons: List[str]) -> Tuple[str, str]:
    key_reasons: List[str] = []
    missing_words: List[str] = []
    ignore_words: List[str] = [
        "and",
        "or",
        "of",
        "in",
        "to",
        "too",
        "at",
        "when",
        "with",
        "for",
        "her",
        "him",
        "while",
    ]

    # loop through all the lines of text in the key reasons section
    # goal is to parse out just the bullet points from this section
    # skip over the first line since it does not contain any keywords
    full_text = ""
    for line in list_reasons[1:]:
        # check to see if the current line contains a period
        if "." in line:
            # consider this the end of the bullet points in the section
            # bullet points are usually followed by some text sentences
            break

        # add the current line to the full text variable
        full_text += line + "\n"

        # split the reason into separate words
        if "," in line:
            list_words = []
            list_items = line.split(",")
            for item in list_items:
                if " " in item:
                    list_words.extend(item.split(" "))
                else:
                    list_words.append(item)
        elif "/" in line:
            list_words = []
            list_items = line.split("/")
            for item in list_items:
                if " " in item:
                    list_words.extend(item.split(" "))
                else:
                    list_words.append(item)
        else:
            list_words = line.split(" ")

        # loop through all the words in the current line
        for word in list_words:
            # convert word to lower case
            word = word.lower().strip()
            if word.startswith("and"):
                word.replace("and", "")
                word = word.strip()

            if len(word) > 1:
                # check to see if the word is in the dictionary of keywords
                if word in dict_keywords:
                    # add word to the list and increment the count
                    key_reasons.append(word)
                    dict_keywords[word] += 1
                else:
                    # check to see if we can ignore the missing word
                    if word not in ignore_words:
                        # add word to the list of missing words from dictionary
                        missing_words.append(word)

    # compute the reasons text string
    reasons_text = "|".join(key_reasons)
    logger.debug("parse_key_reasons: full text length [%s].", len(full_text))
    logger.debug(
        "parse_key_reasons: key words [%s] with length [%s].", len(key_reasons), len(reasons_text)
    )

    # check to see if there were any missing words
    if len(missing_words) > 0:
        missing_text = "|".join(missing_words)
        logger.info("parse_key_reasons: Missing [%s] words [%s].", len(missing_words), missing_text)

    # return the full text of bullet points and the list of keywords found
    return (full_text, reasons_text)


def docx_parse_client(
    file_name: str, document: Document, reasons: str, keywords: str
) -> Optional[DocumentSchema]:
    # ensure the document has at least one table
    if len(document.tables) == 0:
        logger.error("docx_parse_client: [%s] does not have any tables.", file_name)
        return None

    # initialize values used to parse the first table in the document
    table = document.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    if (num_rows != 4) and (num_cols != 2):
        logger.error(
            "docx_parse_client: [%s] invalid table size [%s][%s].", file_name, num_rows, num_cols
        )
        return None

    # initialize the document schema values
    doc_vector = embed_text(reasons)
    doc_uuid = str(uuid.uuid4())
    doc_data = DocumentSchema(
        assessment_uuid=doc_uuid,
        assessment_file=file_name,
        assessment_reasons=reasons,
        assessment_keywords=keywords,
        vector=doc_vector,
    )

    # get list of the unique strings from each cell in the table
    list_info = []
    for r in range(num_rows):
        for c in range(num_cols):
            # parse out the text for the current cell
            text = table.cell(r, c).text.strip()
            text = text.replace("\t", "")
            text = text.replace("\n", "")
            text = text.replace("\u2019", "'")
            text = text.replace("\u2013", "-")

            # ensure the text is not empty
            item_text = text.strip()
            if len(item_text) > 0:
                # if the string contains a comma, only take the first part of it
                if "," in item_text:
                    item_text = item_text.split(",")[0]

                # ensure the string is not a duplicate
                if item_text not in list_info:
                    list_info.append(item_text)

    # loop through the list of strings
    for item in list_info:
        # ensure the required delimiter is found
        delimiter = ":"
        if ":" not in item:
            logger.warning("docx_parse_client: [%s] has invalid item [%s].", file_name, item)
            continue

        # split the string on the colon and get the key and value
        item_split = item.split(":", 1)
        item_key = item_split[0].strip()
        item_value = item_split[1].strip()

        # update the output value based on the key value
        if item_key.startswith("Date"):
            # ensure we have a valid date value
            if len(item_value) == 0:
                logger.warning(
                    "docx_parse_client: [%s] Missing date value [%s][%s].",
                    file_name,
                    item_key,
                    item_value,
                )
                continue

            # ensure date is in correct format
            delimiter = "-"
            if "/" in item_value:
                delimiter = "/"

            # ensure date is in correct format
            date_info = item_value.split(delimiter)
            if len(date_info) == 3:
                if len(date_info[2]) > 4:
                    date_info[2] = date_info[2][:4]
                date_year = int(date_info[2])
                date_month = int(date_info[0])
                date_day = int(date_info[1])
            elif len(date_info) == 2:
                if len(date_info[1]) > 4:
                    date_info[2] = date_info[1][:4]
                date_year = int(date_info[1][:4])
                date_month = int(date_info[0])
                date_day = 1
            else:
                logger.warning(
                    "docx_parse_client: Unable to parse date value [%s][%s].", item_key, item_value
                )
                continue

            # ensure year value is correct
            if date_year < 1900:
                date_year += 2000

            # compute the date value
            date_value = datetime(date_year, date_month, date_day)
            if "Birth" not in item_key:
                # update the assessment date
                doc_data.assessment_date = datetime.strftime(date_value, "%Y-%m-%d")
            else:
                # update the client birth date and compute age value
                doc_data.client_dob = datetime.strftime(date_value, "%Y-%m-%d")
                doc_data.client_age = get_age(
                    str(doc_data.assessment_date), str(doc_data.client_dob)
                )

        elif item_key.startswith("Client"):
            doc_data.client_name = item_value
        elif item_key.startswith("Grade"):
            doc_data.client_grade = get_grade(item_value)
        elif item_key.startswith("From"):
            item_splits = item_value.split(" ")
            item_initials = [item[0] for item in item_splits if len(item) > 0]
            doc_data.assessment_author = item_value
            doc_data.assessment_initials = "".join(item_initials)

    return doc_data


def docx_parse_sections(document: Document, sections_defn: Dict) -> Dict:
    # determine the section delimiter values
    sections_text = [item["text"] for item in sections_defn]

    # loop through all the document paragraphs
    doc_sections = {}
    section_index: int = 0
    section_data: Optional[List] = None
    section_info: Optional[Dict[str, Any]] = None
    for idx, paragraph in enumerate(document.paragraphs):
        # remove extra spaces
        # also ensure unicode characters and others are replaced
        text = paragraph.text.strip()
        text = text.replace("\t", "")
        text = text.replace("\n", "")
        text = text.replace("\u2019", "'")
        text = text.replace("\u2013", "-")

        # ignore empty text lines
        if len(text) > 0:
            # determine if text indicates the start of a new section
            new_index = -1
            if (section_info is not None) and (section_info["title"] == "key-reasons"):
                # only end this section when the next section is found
                # this section may contain other section key words
                if text.startswith("Summary and Impact of Challenges"):
                    new_index = is_section_text(text, sections_text)
            else:
                new_index = is_section_text(text, sections_text)

            # check to see if we found a new section
            if new_index >= 0:
                # check to see if we have an active section
                if (
                    (section_data is not None)
                    and (len(section_data) > 0)
                    and (section_info is not None)
                ):
                    # persist the current section data
                    if not section_info["skip"]:
                        # check to see if the section data should be filtered
                        if section_info["text_range"] is not None:
                            slice_ints = section_info["text_range"].split(":")
                            start_idx = int(slice_ints[0])
                            if (len(slice_ints) == 2) and ((len(slice_ints[1]) > 0)):
                                end_idx = int(slice_ints[1])
                                section_data = section_data[start_idx:end_idx]
                            else:
                                section_data = section_data[start_idx:]

                        # save out the section data
                        section_length = 0
                        for item in section_data:
                            section_length += len(item)
                        doc_sections[section_info["title"]] = section_data
                        # print(section_info["title"], len(section_data), section_length)

                # initialize the new section data
                section_index = new_index
                section_data = []
                section_info = sections_defn[section_index]

                # check to see if we have reached the end of the sections
                if section_index >= len(sections_text):
                    section_data = None
                    break

            elif (section_data is not None) and (section_info is not None):
                if section_info["max_length"] < 0:
                    # add the text to the current section
                    section_data.append(text)
                else:
                    # only add the text if it is less than the max length value
                    if len(text) <= section_info["max_length"]:
                        section_data.append(text)

    return doc_sections


def docx_parse_tables(file_name: str, tables_defn: Dict) -> Optional[Dict]:
    return None
